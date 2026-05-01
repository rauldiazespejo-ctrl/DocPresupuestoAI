import pdfplumber
import docx
import openpyxl
import json
import re
from pathlib import Path

def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

def extract_text_from_docx(file_path: str) -> str:
    doc = docx.Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + "\t"
            text += "\n"
    return text

def extract_text_from_excel(file_path: str) -> str:
    wb = openpyxl.load_workbook(file_path, data_only=True)
    text = ""
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        text += f"\n=== HOJA: {sheet} ===\n"
        for row in ws.iter_rows(values_only=True):
            row_text = "\t".join([str(cell) if cell is not None else "" for cell in row])
            if row_text.strip():
                text += row_text + "\n"
    return text

def extract_text_from_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def extract_text(file_path: str) -> str:
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    elif ext in [".xlsx", ".xls"]:
        return extract_text_from_excel(file_path)
    elif ext in [".txt", ".md"]:
        return extract_text_from_txt(file_path)
    else:
        return ""

def parse_money(text: str) -> float:
    text = text.replace(".", "").replace(",", ".").replace("$", "").strip()
    try:
        return float(text)
    except:
        return 0.0

def find_amounts_in_text(text: str) -> list:
    pattern = r'\$?\s*[\d.,]+(?:\s*(?:UF|CLP|USD|EUR|UTM))?' 
    matches = re.findall(pattern, text)
    return matches
